"""
Context Ingestion - Extract text from various sources for Project context
Handles files, URLs, text, and images with OCR support
"""

import os
import io
import base64
from pathlib import Path
from typing import Dict, Any, Optional, Callable, BinaryIO
from dataclasses import dataclass
import logging
import mimetypes

from PIL import Image

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    pytesseract = None

logger = logging.getLogger(__name__)


@dataclass
class IngestionResult:
    """Result of context ingestion"""
    success: bool
    content: str
    metadata: Dict[str, Any]
    error: Optional[str] = None


class ContextIngestionEngine:
    """Engine for ingesting various content types into Project context"""

    # Supported file extensions by category
    TEXT_EXTENSIONS = {'.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml', '.csv'}
    CODE_EXTENSIONS = {'.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.cpp', '.c', '.h', '.hpp', '.rs', '.go', '.rb', '.php', '.swift', '.kt', '.scala', '.r', '.m', '.mm'}
    DOC_EXTENSIONS = {'.pdf', '.docx', '.doc'}
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'}

    def __init__(self):
        self.url_processor = None
        self._init_url_processor()

    def _init_url_processor(self):
        """Initialize URL processor if available"""
        try:
            from src.url_processor import URLProcessor
            self.url_processor = URLProcessor()
        except ImportError as e:
            logger.warning(f"URLProcessor not available: {e}")

    def ingest_file(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> IngestionResult:
        """Ingest a file and extract its text content"""
        path = Path(file_path)

        if not path.exists():
            return IngestionResult(
                success=False,
                content="",
                metadata={},
                error=f"File not found: {file_path}"
            )

        extension = path.suffix.lower()

        try:
            if extension in self.IMAGE_EXTENSIONS:
                return self._ingest_image(file_path, progress_callback)
            elif extension == '.pdf':
                return self._ingest_pdf(file_path, progress_callback)
            elif extension in {'.docx', '.doc'}:
                return self._ingest_word(file_path, progress_callback)
            elif extension in self.TEXT_EXTENSIONS or extension in self.CODE_EXTENSIONS:
                return self._ingest_text_file(file_path, progress_callback)
            else:
                # Try as text file
                return self._ingest_text_file(file_path, progress_callback)

        except Exception as e:
            logger.error(f"Error ingesting file {file_path}: {e}")
            return IngestionResult(
                success=False,
                content="",
                metadata={"file_path": file_path, "extension": extension},
                error=str(e)
            )

    def ingest_url(
        self,
        url: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> IngestionResult:
        """Ingest content from a URL"""
        if self.url_processor is None:
            return IngestionResult(
                success=False,
                content="",
                metadata={"url": url},
                error="URL processor not available"
            )

        try:
            if progress_callback:
                progress_callback(f"🌐 Processing URL: {url}")

            result = self.url_processor.process(url, progress_callback)

            return IngestionResult(
                success=True,
                content=result.get("content", ""),
                metadata={
                    "url": url,
                    "title": result.get("title", ""),
                    "archive_url": result.get("archive_url", ""),
                    "word_count": result.get("metadata", {}).get("word_count", 0),
                    "content_length": result.get("metadata", {}).get("content_length", 0),
                }
            )

        except Exception as e:
            logger.error(f"Error ingesting URL {url}: {e}")
            return IngestionResult(
                success=False,
                content="",
                metadata={"url": url},
                error=str(e)
            )

    def ingest_text(
        self,
        text: str,
        title: str = "Text Note"
    ) -> IngestionResult:
        """Ingest plain text content"""
        return IngestionResult(
            success=True,
            content=text,
            metadata={
                "title": title,
                "char_count": len(text),
                "word_count": len(text.split()),
                "line_count": len(text.splitlines()),
            }
        )

    def ingest_image(
        self,
        image_path: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> IngestionResult:
        """Ingest an image and extract text via OCR"""
        return self._ingest_image(image_path, progress_callback)

    def _ingest_text_file(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> IngestionResult:
        """Ingest a plain text or code file"""
        path = Path(file_path)

        if progress_callback:
            progress_callback(f"📄 Reading file: {path.name}")

        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        content = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            return IngestionResult(
                success=False,
                content="",
                metadata={"file_path": file_path},
                error="Could not decode file with any supported encoding"
            )

        # Detect language/type based on extension
        extension = path.suffix.lower()
        file_type = "code" if extension in self.CODE_EXTENSIONS else "text"

        return IngestionResult(
            success=True,
            content=content,
            metadata={
                "file_path": file_path,
                "filename": path.name,
                "extension": extension,
                "file_type": file_type,
                "char_count": len(content),
                "line_count": len(content.splitlines()),
            }
        )

    def _ingest_pdf(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> IngestionResult:
        """Ingest a PDF file"""
        try:
            import PyPDF2

            if progress_callback:
                progress_callback(f"📑 Extracting PDF: {Path(file_path).name}")

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)

                content_parts = []
                for i, page in enumerate(pdf_reader.pages):
                    if progress_callback and i % 5 == 0:
                        progress_callback(f"📑 Processing page {i+1}/{num_pages}")
                    content_parts.append(page.extract_text())

                content = "\n\n".join(content_parts)

            return IngestionResult(
                success=True,
                content=content,
                metadata={
                    "file_path": file_path,
                    "filename": Path(file_path).name,
                    "file_type": "pdf",
                    "page_count": num_pages,
                    "char_count": len(content),
                }
            )

        except ImportError:
            return IngestionResult(
                success=False,
                content="",
                metadata={"file_path": file_path},
                error="PyPDF2 not installed. Install with: pip install PyPDF2"
            )
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return IngestionResult(
                success=False,
                content="",
                metadata={"file_path": file_path},
                error=str(e)
            )

    def _ingest_word(
        self,
        file_path: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> IngestionResult:
        """Ingest a Word document (.docx or .doc)"""
        try:
            from docx import Document

            if progress_callback:
                progress_callback(f"📝 Extracting Word doc: {Path(file_path).name}")

            doc = Document(file_path)

            content_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)

            content = "\n\n".join(content_parts)

            return IngestionResult(
                success=True,
                content=content,
                metadata={
                    "file_path": file_path,
                    "filename": Path(file_path).name,
                    "file_type": "word",
                    "paragraph_count": len(doc.paragraphs),
                    "char_count": len(content),
                }
            )

        except ImportError:
            return IngestionResult(
                success=False,
                content="",
                metadata={"file_path": file_path},
                error="python-docx not installed. Install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Error processing Word doc {file_path}: {e}")
            return IngestionResult(
                success=False,
                content="",
                metadata={"file_path": file_path},
                error=str(e)
            )

    def _ingest_image(
        self,
        image_path: str,
        progress_callback: Optional[Callable[[str], None]] = None
    ) -> IngestionResult:
        """Ingest an image and extract text via OCR"""
        if not PYTESSERACT_AVAILABLE:
            return IngestionResult(
                success=False,
                content="",
                metadata={"file_path": image_path},
                error="pytesseract not installed. Install with: pip install pytesseract"
            )

        try:
            if progress_callback:
                progress_callback(f"🖼️ Processing image: {Path(image_path).name}")

            # Open image
            image = Image.open(image_path)

            # Get image metadata
            metadata = {
                "file_path": image_path,
                "filename": Path(image_path).name,
                "file_type": "image",
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
            }

            if progress_callback:
                progress_callback("🔍 Running OCR...")

            # Perform OCR
            text = pytesseract.image_to_string(image)

            metadata["char_count"] = len(text)
            metadata["word_count"] = len(text.split()) if text else 0

            return IngestionResult(
                success=True,
                content=text,
                metadata=metadata
            )

        except Exception as e:
            logger.error(f"Error processing image {image_path}: {e}")
            return IngestionResult(
                success=False,
                content="",
                metadata={"file_path": image_path},
                error=str(e)
            )

    def get_file_type(self, file_path: str) -> str:
        """Get the type of file based on extension"""
        extension = Path(file_path).suffix.lower()

        if extension in self.IMAGE_EXTENSIONS:
            return "image"
        elif extension == '.pdf':
            return "pdf"
        elif extension in {'.docx', '.doc'}:
            return "word"
        elif extension in self.CODE_EXTENSIONS:
            return "code"
        elif extension in self.TEXT_EXTENSIONS:
            return "text"
        else:
            return "unknown"

    def is_supported_file(self, file_path: str) -> bool:
        """Check if a file type is supported"""
        extension = Path(file_path).suffix.lower()
        all_supported = (
            self.TEXT_EXTENSIONS |
            self.CODE_EXTENSIONS |
            self.DOC_EXTENSIONS |
            self.IMAGE_EXTENSIONS
        )
        return extension in all_supported


# Global instance for reuse
_ingestion_engine = None


def get_ingestion_engine() -> ContextIngestionEngine:
    """Get or create the global ingestion engine"""
    global _ingestion_engine
    if _ingestion_engine is None:
        _ingestion_engine = ContextIngestionEngine()
    return _ingestion_engine
